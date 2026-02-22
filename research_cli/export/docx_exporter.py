"""
docx_exporter.py — Convert markdown thesis draft to USIL-formatted .docx.

Formatting:
    - Font: Times New Roman, 12pt
    - Line spacing: 1.5
    - Margins: 2.54 cm (1 inch) all sides — USIL standard
    - Headings: Heading 1/2/3 mapped from markdown #/##/###
    - Tables: parsed from markdown pipe-delimited tables
    - Bold / italic: parsed from **bold** and *italic* markers
    - Section separators (---) are consumed silently
"""

import re

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT


# ---------------------------------------------------------------------------
# Style constants
# ---------------------------------------------------------------------------

FONT_NAME = "Times New Roman"
FONT_SIZE_PT = 12
LINE_SPACING = 1.5
MARGIN_CM = 2.54

HEADING_SIZES = {1: 16, 2: 14, 3: 12}


# ---------------------------------------------------------------------------
# Document setup
# ---------------------------------------------------------------------------

def _setup_styles(doc: Document) -> None:
    """Configure default document styles for USIL formatting."""
    # Default paragraph style
    style = doc.styles["Normal"]
    font = style.font
    font.name = FONT_NAME
    font.size = Pt(FONT_SIZE_PT)
    pf = style.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = LINE_SPACING
    pf.space_after = Pt(0)
    pf.space_before = Pt(0)

    # Heading styles
    for level in (1, 2, 3):
        style_name = f"Heading {level}"
        if style_name in doc.styles:
            h_style = doc.styles[style_name]
            h_font = h_style.font
            h_font.name = FONT_NAME
            h_font.size = Pt(HEADING_SIZES[level])
            h_font.bold = True
            h_font.color.rgb = RGBColor(0, 0, 0)
            h_pf = h_style.paragraph_format
            h_pf.space_before = Pt(12)
            h_pf.space_after = Pt(6)
            h_pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
            h_pf.line_spacing = LINE_SPACING


def _setup_margins(doc: Document) -> None:
    """Set USIL standard margins (2.54 cm / 1 inch all sides)."""
    for section in doc.sections:
        section.top_margin = Cm(MARGIN_CM)
        section.bottom_margin = Cm(MARGIN_CM)
        section.left_margin = Cm(MARGIN_CM)
        section.right_margin = Cm(MARGIN_CM)


# ---------------------------------------------------------------------------
# Inline formatting (bold / italic)
# ---------------------------------------------------------------------------

_INLINE_RE = re.compile(
    r"(\*\*\*(.+?)\*\*\*)"   # ***bold italic***
    r"|(\*\*(.+?)\*\*)"       # **bold**
    r"|(\*(.+?)\*)"           # *italic*
)


def _add_formatted_text(paragraph, text: str) -> None:
    """Add text to a paragraph, parsing **bold**, *italic*, and ***both***."""
    pos = 0
    for m in _INLINE_RE.finditer(text):
        # Add plain text before this match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])

        if m.group(2):  # ***bold italic***
            run = paragraph.add_run(m.group(2))
            run.bold = True
            run.italic = True
        elif m.group(4):  # **bold**
            run = paragraph.add_run(m.group(4))
            run.bold = True
        elif m.group(6):  # *italic*
            run = paragraph.add_run(m.group(6))
            run.italic = True

        pos = m.end()

    # Remaining plain text
    if pos < len(text):
        paragraph.add_run(text[pos:])


def _add_paragraph(doc: Document, text: str) -> None:
    """Add a paragraph with inline formatting applied."""
    p = doc.add_paragraph()
    _add_formatted_text(p, text)
    # Apply font to all runs
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)


# ---------------------------------------------------------------------------
# Table parsing
# ---------------------------------------------------------------------------

def _is_table_row(line: str) -> bool:
    """Check if a line looks like a markdown table row."""
    stripped = line.strip()
    return stripped.startswith("|") and stripped.endswith("|")


def _is_separator_row(line: str) -> bool:
    """Check if a line is a table separator (|---|---|)."""
    stripped = line.strip()
    return bool(re.match(r"^\|[\s\-:|]+\|$", stripped))


def _parse_table_row(line: str) -> list[str]:
    """Extract cell contents from a markdown table row."""
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def _add_table(doc: Document, rows: list[list[str]]) -> None:
    """Add a formatted table to the document."""
    if not rows:
        return

    num_cols = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    for i, row_data in enumerate(rows):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            if j < num_cols:
                cell = row.cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                _add_formatted_text(p, cell_text)
                # Apply font
                for run in p.runs:
                    run.font.name = FONT_NAME
                    run.font.size = Pt(FONT_SIZE_PT)
                    if i == 0:  # Header row bold
                        run.bold = True


# ---------------------------------------------------------------------------
# List parsing
# ---------------------------------------------------------------------------

def _is_list_item(line: str) -> bool:
    """Check if a line is a numbered or bulleted list item."""
    stripped = line.strip()
    return bool(
        re.match(r"^\d+\.\s", stripped)
        or stripped.startswith("- ")
        or stripped.startswith("* ")
    )


def _get_list_text(line: str) -> str:
    """Extract text content from a list item line."""
    stripped = line.strip()
    # Numbered: "1. text"
    m = re.match(r"^\d+\.\s+(.*)", stripped)
    if m:
        return m.group(1)
    # Bulleted: "- text" or "* text"
    if stripped.startswith("- ") or stripped.startswith("* "):
        return stripped[2:]
    return stripped


def _add_list_item(doc: Document, text: str, numbered: bool = False) -> None:
    """Add a list item paragraph."""
    style = "List Number" if numbered else "List Bullet"
    # Fallback if style doesn't exist
    try:
        p = doc.add_paragraph(style=style)
    except KeyError:
        p = doc.add_paragraph()
    _add_formatted_text(p, text)
    for run in p.runs:
        run.font.name = FONT_NAME
        run.font.size = Pt(FONT_SIZE_PT)


# ---------------------------------------------------------------------------
# Main converter
# ---------------------------------------------------------------------------

def _parse_heading(line: str) -> tuple[int, str] | None:
    """Parse a markdown heading line. Returns (level, text) or None."""
    m = re.match(r"^(#{1,3})\s+(.*)", line.strip())
    if m:
        return len(m.group(1)), m.group(2).strip()
    return None


def markdown_to_docx(markdown_text: str, output_path: str) -> str:
    """Convert a markdown thesis draft to a USIL-formatted .docx file.

    Parameters
    ----------
    markdown_text : str
        The complete thesis draft in markdown format.
    output_path : str
        Path where the .docx file will be saved.

    Returns
    -------
    str
        The output path of the saved .docx file.
    """
    doc = Document()
    _setup_styles(doc)
    _setup_margins(doc)

    lines = markdown_text.split("\n")
    i = 0

    while i < len(lines):
        line = lines[i]
        stripped = line.strip()

        # Skip empty lines
        if not stripped:
            i += 1
            continue

        # Skip horizontal rules (section separators)
        if re.match(r"^-{3,}$", stripped) or re.match(r"^\*{3,}$", stripped):
            i += 1
            continue

        # Heading
        heading = _parse_heading(line)
        if heading:
            level, text = heading
            doc.add_heading(text, level=level)
            i += 1
            continue

        # Table block
        if _is_table_row(stripped):
            table_rows = []
            while i < len(lines) and _is_table_row(lines[i].strip()):
                if not _is_separator_row(lines[i]):
                    table_rows.append(_parse_table_row(lines[i]))
                i += 1
            _add_table(doc, table_rows)
            continue

        # List item
        if _is_list_item(stripped):
            numbered = bool(re.match(r"^\d+\.\s", stripped))
            text = _get_list_text(stripped)
            _add_list_item(doc, text, numbered=numbered)
            i += 1
            continue

        # Regular paragraph — collect consecutive non-blank, non-special lines
        para_lines = []
        while i < len(lines):
            current = lines[i].strip()
            if not current:
                break
            if _parse_heading(lines[i]):
                break
            if _is_table_row(current):
                break
            if re.match(r"^-{3,}$", current) or re.match(r"^\*{3,}$", current):
                break
            if _is_list_item(current) and not para_lines:
                break
            # If we already have paragraph text and hit a list item, break
            if _is_list_item(current) and para_lines:
                break
            para_lines.append(current)
            i += 1

        if para_lines:
            _add_paragraph(doc, " ".join(para_lines))

    doc.save(output_path)
    return output_path
